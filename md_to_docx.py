#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将《项目可行性分析报告.md》转换为 Word (.docx) 文档
保留：标题层级、表格、代码块、列表、加粗、引用块
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"d:\A-Users\Desktop\zhongheshixun\项目可行性分析报告.md"
DST = r"d:\A-Users\Desktop\zhongheshixun\项目可行性分析报告.docx"

# ------------------- 工具函数 -------------------

def set_cn_font(run, name="宋体", size=10.5, bold=False, color=None):
    """设置中英文字体（中文宋体 / 西文 Times New Roman）"""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), 'Times New Roman')
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, color_hex):
    """设置表格单元格底色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_table_borders(table):
    """为表格添加边框"""
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_inline_runs(paragraph, text):
    """处理 **加粗** 与 `代码` 行内格式"""
    # 拆分文本为 (text, bold, mono) 三元组
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            seg = text[pos:m.start()]
            r = paragraph.add_run(seg)
            set_cn_font(r)
        token = m.group()
        if token.startswith('**'):
            r = paragraph.add_run(token[2:-2])
            set_cn_font(r, bold=True)
        else:  # `code`
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            rPr = r._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Consolas')
            rFonts.set(qn('w:hAnsi'), 'Consolas')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rPr.append(rFonts)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_cn_font(r)

# ------------------- 解析 markdown -------------------

def parse_markdown(md_text):
    """返回 block 列表，每个 block 是 dict：
       {type: 'h1'/'h2'/'h3'/'h4'/'p'/'code'/'table'/'ul'/'ol'/'hr'/'quote', content/rows}
    """
    lines = md_text.split('\n')
    blocks = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        # 跳过空行
        if not stripped.strip():
            i += 1
            continue

        # 水平线
        if re.match(r'^-{3,}$', stripped):
            blocks.append({'type': 'hr'})
            i += 1
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            blocks.append({'type': f'h{level}', 'text': m.group(2).strip()})
            i += 1
            continue

        # 引用块
        if stripped.startswith('> '):
            quote_lines = []
            while i < n and lines[i].startswith('> '):
                quote_lines.append(lines[i][2:].strip())
                i += 1
            blocks.append({'type': 'quote', 'text': '\n'.join(quote_lines)})
            continue

        # 代码块（围栏）
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].rstrip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过闭合 ```
            blocks.append({'type': 'code', 'lang': lang, 'text': '\n'.join(code_lines)})
            continue

        # 表格
        if '|' in stripped and i + 1 < n and re.match(r'^\s*\|?\s*:?-+:?(\s*\|\s*:?-+:?)+\s*\|?\s*$', lines[i+1]):
            rows = []
            while i < n and '|' in lines[i] and lines[i].strip():
                # 解析一行
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(row)
                i += 1
            blocks.append({'type': 'table', 'rows': rows})
            continue

        # 无序列表
        if re.match(r'^[\-\*]\s+', stripped):
            items = []
            while i < n and re.match(r'^[\-\*]\s+', lines[i].rstrip()):
                items.append(re.sub(r'^[\-\*]\s+', '', lines[i].rstrip()))
                i += 1
            blocks.append({'type': 'ul', 'items': items})
            continue

        # 有序列表
        if re.match(r'^\d+\.\s+', stripped):
            items = []
            while i < n and re.match(r'^\d+\.\s+', lines[i].rstrip()):
                items.append(re.sub(r'^\d+\.\s+', '', lines[i].rstrip()))
                i += 1
            blocks.append({'type': 'ol', 'items': items})
            continue

        # 普通段落（可能跨行）
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(r'^(#{1,6}\s|>\s|```|\||[\-\*]\s|\d+\.\s|-{3,})', lines[i].rstrip()):
            para_lines.append(lines[i].rstrip())
            i += 1
        blocks.append({'type': 'p', 'text': ' '.join(para_lines)})

    return blocks

# ------------------- 写入 docx -------------------

def write_docx(blocks, out_path):
    doc = Document()

    # 全局样式：默认正文
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), '宋体')

    # 页边距
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 页脚：页码
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
    set_cn_font(run, size=9)

    # 标题样式
    h_styles = {
        'h1': ('Heading 1', 18, True, RGBColor(0x1F, 0x4E, 0x79)),
        'h2': ('Heading 2', 15, True, RGBColor(0x2E, 0x74, 0xB5)),
        'h3': ('Heading 3', 13, True, RGBColor(0x1F, 0x4E, 0x79)),
        'h4': ('Heading 4', 11, True, RGBColor(0x2E, 0x74, 0xB5)),
        'h5': ('Heading 5', 11, True, RGBColor(0x2E, 0x74, 0xB5)),
        'h6': ('Heading 6', 10.5, True, RGBColor(0x2E, 0x74, 0xB5)),
    }

    for b in blocks:
        t = b['type']
        if t in h_styles:
            sty_name, size, bold, color = h_styles[t]
            p = doc.add_paragraph(style=sty_name)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(b['text'])
            set_cn_font(r, size=size, bold=bold, color=tuple(int(str(color)[i:i+2], 16) for i in (0, 2, 4)))
        elif t == 'p':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Pt(21)  # 首行缩进
            add_inline_runs(p, b['text'])
        elif t == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for line in b['text'].split('\n'):
                if line.startswith('**') and '：' in line and not p.runs:
                    # 加粗开头的引用
                    add_inline_runs(p, line)
                else:
                    if p.runs:
                        p.add_run('\n')
                    add_inline_runs(p, line)
            # 左侧竖线
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left'); left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24'); left.set(qn('w:color'), '2E74B5')
            pBdr.append(left)
            pPr.append(pBdr)
        elif t == 'code':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            # 底色
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F2F2F2')
            pPr.append(shd)
            for line_idx, line in enumerate(b['text'].split('\n')):
                if line_idx > 0:
                    br = OxmlElement('w:br')
                    p.runs[-1]._r.append(br) if p.runs else p.add_run()._r.append(br)
                r = p.add_run(line if line else ' ')
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                rPr = r._element.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Consolas'); rFonts.set(qn('w:hAnsi'), 'Consolas'); rFonts.set(qn('w:eastAsia'), '宋体')
                rPr.append(rFonts)
        elif t == 'ul':
            for item in b['items']:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                add_inline_runs(p, item)
        elif t == 'ol':
            for item in b['items']:
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                add_inline_runs(p, item)
        elif t == 'hr':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:color'), '808080')
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif t == 'table':
            rows = b['rows']
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            # 规范化
            rows = [r + [''] * (ncols - len(r)) for r in rows]
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            set_table_borders(table)
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # 清空默认段落
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.2
                    add_inline_runs(p, cell_text)
                    if r_idx == 0:
                        # 表头加粗 + 底色
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        shade_cell(cell, '2E74B5')
                    else:
                        # 隔行底色
                        if r_idx % 2 == 0:
                            shade_cell(cell, 'F2F7FB')
                    # 字体设置
                    for run in p.runs:
                        if not run.font.name:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9.5)
                        set_cn_font(run, size=9.5, bold=(r_idx == 0))

    doc.save(out_path)
    print(f"OK: {out_path}")

if __name__ == '__main__':
    with open(SRC, 'r', encoding='utf-8') as f:
        md = f.read()
    blocks = parse_markdown(md)
    print(f"解析到 {len(blocks)} 个 block")
    write_docx(blocks, DST)
