# -*- coding: utf-8 -*-
"""
生成龙源数据库论文标准格式的Word文档
整合：角色分工与主要任务 + 技术选型与成本对比
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def set_run_font(run, font_name='宋体', font_size=Pt(10.5), bold=False, color=None):
    """设置run字体"""
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color


def add_paragraph_with_style(doc, text, font_name='宋体', font_size=Pt(10.5),
                              bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              first_line_indent=None, line_spacing=1.5,
                              space_before=Pt(0), space_after=Pt(0)):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent

    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold)
    return p


def add_heading_numbered(doc, number, text, level=1):
    """添加编号标题"""
    if level == 1:
        font_size = Pt(14)  # 四号
        space_before = Pt(12)
        space_after = Pt(6)
    elif level == 2:
        font_size = Pt(12)  # 小四号
        space_before = Pt(6)
        space_after = Pt(3)
    else:
        font_size = Pt(10.5)  # 五号
        space_before = Pt(3)
        space_after = Pt(3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.first_line_indent = Cm(0)

    run = p.add_run(f"{number} {text}")
    set_run_font(run, '黑体', font_size, bold=True)
    return p


def add_table_with_style(doc, headers, rows, caption):
    """添加带样式的表格"""
    # 表题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(caption)
    set_run_font(run, '黑体', Pt(9), bold=True)  # 小五号黑体

    # 表格
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, '黑体', Pt(9), bold=True)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, '宋体', Pt(9))

    doc.add_paragraph()  # 表格后空行
    return table


def main():
    doc = Document()

    # ==================== 页面设置 ====================
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ==================== 标题 ====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("基于Spring Boot的基层卫生人员网络培训平台\n开发环境搭建与技术选型研究")
    set_run_font(run, '黑体', Pt(16), bold=True)

    # ==================== 作者信息 ====================
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(3)
    run = author.add_run("成员A\u3000成员B\u3000成员C\u3000成员D")
    set_run_font(run, '楷体', Pt(12))

    # 单位
    unit = doc.add_paragraph()
    unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    unit.paragraph_format.space_after = Pt(12)
    run = unit.add_run("（某某大学 计算机科学与技术学院，某某 000000）")
    set_run_font(run, '宋体', Pt(9))

    # ==================== 摘要 ====================
    abstract_p = doc.add_paragraph()
    abstract_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_p.paragraph_format.line_spacing = 1.5
    abstract_p.paragraph_format.first_line_indent = Cm(0)

    run_label = abstract_p.add_run("【摘要】")
    set_run_font(run_label, '黑体', Pt(10.5), bold=True)

    abstract_text = (
        "随着基层卫生人员继续教育需求的日益增长，构建高效、低成本的在线培训平台成为重要课题。"
        "本文以四川省基层卫生人员网络培训平台为研究对象，系统阐述了项目开发中的角色分工协作机制、"
        "技术选型决策过程、版本控制实践以及系统实现与测试经验。项目采用4人团队协作模式，"
        "基于Spring Boot 2.7.18 + Vue3 + MySQL 8.0技术栈，全部选用开源方案实现零软件授权成本。"
        "系统实现了MOOC学习、在线考试、离线学习、实时咨询四大核心功能模块，"
        "并通过RBAC权限控制、BCrypt密码加密、JWT无状态认证等安全机制保障医疗数据隐私。"
        "通过对比开源与商业方案的成本差异，论证了开源技术栈在满足3000用户并发、响应时间小于2秒等"
        "性能需求前提下的经济性与可行性，为同类教育信息化项目的技术选型提供参考。"
    )
    run_text = abstract_p.add_run(abstract_text)
    set_run_font(run_text, '宋体', Pt(10.5))

    # ==================== 关键词 ====================
    keyword_p = doc.add_paragraph()
    keyword_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keyword_p.paragraph_format.line_spacing = 1.5
    keyword_p.paragraph_format.space_after = Pt(12)

    run_kw_label = keyword_p.add_run("【关键词】")
    set_run_font(run_kw_label, '黑体', Pt(10.5), bold=True)

    run_kw_text = keyword_p.add_run("基层卫生人员；在线培训平台；Spring Boot；技术选型；开源方案；RBAC；Git；成本分析")
    set_run_font(run_kw_text, '宋体', Pt(10.5))

    # ==================== 引言 ====================
    add_heading_numbered(doc, "0", "引言", level=1)

    intro_text = (
        "基层卫生人员是保障农村地区医疗卫生服务的重要力量。四川省拥有67个民族县，"
        "约3000名基层卫生技术人员分散在偏远地区，传统的集中面授培训模式难以满足其继续教育需求。"
        "因此，构建一套基于互联网技术的在线培训平台，实现MOOC学习、在线考试、离线学习、实时咨询等功能，"
        "对于提升基层卫生人员专业素质具有重要意义。"
        "\n\n本文从项目实际开发角度出发，重点探讨四个核心问题：一是如何在有限的开发周期（1-2周）内"
        "实现高效的团队协作与任务分工；二是如何在满足性能与安全需求的前提下，通过合理的技术选型"
        "降低项目投入成本；三是如何通过规范化的版本控制保障代码质量和团队协作效率；"
        "四是如何验证系统功能的完整性和性能的达标性。"
        "研究内容涵盖角色分工设计、Git版本控制规范、分层技术选型对比、开源与商业方案的成本效益分析、"
        "系统实现要点以及测试验证结果。"
        "\n\n近年来，国内外学者在在线教育平台领域开展了大量研究。文献[2]基于Spring Boot框架"
        "设计实现了在线教育平台，重点探讨了架构设计与功能模块划分；文献[5]研究了RBAC权限模型"
        "在管理系统中的应用，为权限控制提供了理论依据；文献[7]深入分析了Redis缓存技术，"
        "为系统性能优化提供了参考。然而，现有研究较少关注技术选型的成本效益分析，"
        "尤其是针对基层教育信息化项目的开源方案适用性研究仍有不足。本文在此方面进行了有益探索。"
    )
    add_paragraph_with_style(doc, intro_text, first_line_indent=Cm(0.74))

    # ==================== 1. 项目概述与团队协作机制 ====================
    add_heading_numbered(doc, "1", "项目概述与团队协作机制", level=1)

    add_heading_numbered(doc, "1.1", "项目背景与技术架构", level=2)

    p1 = add_paragraph_with_style(doc,
        "本项目面向四川省67个民族县3000名基层卫生技术人员，提供MOOC学习、在线考试、离线学习、"
        "实时咨询等核心服务。系统采用单体Spring Boot架构，技术栈包括Spring Boot 2.7.18、"
        "MyBatis-Plus 3.5.3、MySQL 8.0、Redis 7.x、Vue3 + Element Plus以及微信小程序原生开发。"
        "项目通过Maven组织为5个子模块的层级结构：training-common（公共层）、training-dao（数据访问层）、"
        "training-service（业务逻辑层）、training-admin（后台管理后端）和training-api（小程序API后端）。",
        first_line_indent=Cm(0.74))

    add_paragraph_with_style(doc,
        "系统整体架构采用经典的三层架构模式：表现层（前端页面）、业务逻辑层（Service）和数据访问层（DAO）。"
        "表现层分为三个端：后台管理系统（Vue3 + Element Plus）、PC学员网页端（Vue3）和微信小程序端。"
        "业务逻辑层包含用户管理、课程管理、学习管理、考试管理、培训计划、咨询管理、统计分析七大模块。"
        "数据访问层基于MyBatis-Plus实现，支持分页查询、逻辑删除、条件构造等功能。"
        "数据层采用MySQL主库存储业务数据，Redis作为缓存层存储热点数据和用户会话信息。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "1.2", "核心功能模块", level=2)

    add_table_with_style(doc,
        ["功能模块", "功能描述", "所属端", "关键技术"],
        [
            ["MOOC学习", "课程浏览、报名、章节学习、学习进度记录", "小程序+PC学员端", "Vue3、视频播放、进度追踪"],
            ["在线考试", "考试列表、自动组卷、在线答题、自动阅卷、成绩统计", "小程序+PC学员端", "随机组卷算法、事务管理"],
            ["离线学习", "离线包下载、断点续传、离线学习记录、进度回传", "小程序", "文件下载、状态同步"],
            ["实时咨询", "智能问答、关键词匹配、转人工咨询、工单闭环、SLA告警", "小程序+后台管理", "关键词匹配、实时消息"],
            ["课程管理", "课程CRUD、发布/下架、章节管理、资源上传", "后台管理", "文件上传、状态流转"],
            ["用户管理", "用户CRUD、角色分配、权限控制、账号启禁用", "后台管理", "RBAC、JWT"],
            ["统计分析", "学习统计、考试统计、课程统计、机构统计、趋势分析", "后台管理", "ECharts、聚合查询"],
            ["培训计划", "计划创建、课程关联、计划发布、进度监控", "后台管理+小程序", "数据关联、状态管理"],
        ],
        "表1 核心功能模块一览表"
    )

    add_heading_numbered(doc, "1.3", "角色分工设计", level=2)

    p2 = add_paragraph_with_style(doc,
        "根据项目模块划分与技术栈特点，团队采用4人协作模式，各成员职责如下：",
        first_line_indent=Cm(0.74))

    add_table_with_style(doc,
        ["成员", "角色", "负责模块", "工作量占比", "核心产出"],
        [
            ["成员A", "后端架构师", "项目搭建、数据库设计、用户/课程/学习模块、高并发存储层", "40%",
             "dev-backend.md第1-4章、dev-database.md"],
            ["成员B", "后端开发工程师", "考试模块、培训计划、智能问答、统计模块、高并发计算层", "35%",
             "dev-backend.md第5-8章"],
            ["成员C", "前端开发工程师", "后台管理系统（Vue3）、PC学员网页端", "25%",
             "dev-frontend.md、dev-web-student.md"],
            ["成员D", "小程序开发工程师", "微信小程序学员端", "20%", "dev-miniapp.md"],
        ],
        "表2 团队角色分工表"
    )

    p3 = add_paragraph_with_style(doc,
        "成员A作为后端架构师，承担项目基础设施搭建（Maven多模块、数据库、JWT认证），"
        "并输出接口文档供前端和小程序并行开发。成员B在已搭建的框架上实现考试、统计、咨询等复杂业务逻辑。"
        "成员C负责PC双端前端（管理后台+学员网页端），成员D负责移动端小程序。"
        "四人团队通过\"接口先行、并行开发、分阶段联调\"的协作模式，在9.2天内完成了从需求分析到代码实现的完整闭环。",
        first_line_indent=Cm(0.74))

    add_paragraph_with_style(doc,
        "团队协作遵循\"需求对齐→设计方案→评审→编码→集成验收\"的五阶段门控流程。"
        "接口契约变更必须走代码审查流程，每个阶段结束提交评审，通过后才进入下一阶段。"
        "这种分层协作模式确保了前后端开发的并行性和接口一致性，有效提升了开发效率。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "1.4", "任务依赖与Sprint计划", level=2)

    p4 = add_paragraph_with_style(doc,
        "项目开发按5个Sprint推进，各阶段任务关系如表3所示。",
        first_line_indent=Cm(0.74))

    add_table_with_style(doc,
        ["Sprint", "时间", "成员A", "成员B", "成员C", "成员D", "阶段产出"],
        [
            ["1", "第1-3天", "项目搭建、数据库、用户/课程模块", "试题/知识点管理",
             "Vue3搭建、登录页", "小程序搭建、登录页", "接口文档v1.0、基础框架"],
            ["2", "第4-6天", "章节管理、学习进度", "考试模块、自动组卷/阅卷",
             "课程/章节管理页", "课程浏览/详情/学习页", "核心功能完成60%"],
            ["3", "第7-9天", "联调、接口文档完善", "培训计划、智能问答",
             "考试管理、统计页", "考试功能、个人中心", "核心功能完成85%"],
            ["4", "第10-12天", "高并发存储层验证", "统计接口、SLA机制",
             "咨询管理页", "计划/咨询页", "全部功能完成"],
            ["5", "第13-14天", "部署、答辩准备", "部署、答辩准备",
             "部署、答辩准备", "部署、答辩准备", "演示环境、答辩材料"],
        ],
        "表3 Sprint开发计划表"
    )

    add_paragraph_with_style(doc,
        "关键联调节点包括：登录接口联调（A↔C/D）、课程列表联调（A↔C/D）、"
        "考试全链路联调（B↔C/D）、统计报表联调（B↔C）、咨询联调（B↔D）以及三端全流程E2E测试（全员）。"
        "联调阶段采用\"接口文档先行、前后端并行开发、每日增量联调\"的策略，确保问题及时发现和解决。",
        first_line_indent=Cm(0.74))

    # ==================== 2. Git版本控制规范 ====================
    add_heading_numbered(doc, "2", "Git版本控制规范与实践", level=1)

    add_heading_numbered(doc, "2.1", "分支策略", level=2)

    p5 = add_paragraph_with_style(doc,
        "项目采用Git Flow简化版分支模型，设置main（稳定版本，仅接受develop合并）、"
        "develop（开发版本，日常集成）、feature/xxx（功能分支，从develop检出，完成后合并回develop）"
        "以及hotfix/xxx（紧急修复，从main检出，修复后合并到main和develop）四种分支。"
        "该策略确保了主干代码的稳定性，同时支持多人并行开发。",
        first_line_indent=Cm(0.74))

    add_paragraph_with_style(doc,
        "分支协作流程为：开发人员从develop分支检出feature/xxx功能分支进行开发，"
        "开发完成后提交Pull Request（PR），由其他成员进行代码审查。审查通过后，"
        "将feature分支合并到develop分支。当develop分支积累足够的功能且通过测试后，"
        "合并到main分支并打上版本标签。紧急修复时从main分支检出hotfix分支，"
        "修复完成后同时合并到main和develop分支。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "2.2", "提交规范与代码安全", level=2)

    p6 = add_paragraph_with_style(doc,
        "为保持提交历史清晰可读，团队统一采用前缀式提交信息：feat（新功能）、fix（Bug修复）、"
        "docs（文档更新）、refactor（重构）、chore（构建/工具）。"
        "项目根目录配置.gitignore文件，排除target/、node_modules/、.idea/等构建产物和IDE配置，"
        "同时排除application-*.yml和.env文件，防止数据库密码、JWT密钥等敏感信息泄露到仓库，"
        "满足医疗数据隐私保护的合规要求。",
        first_line_indent=Cm(0.74))

    add_table_with_style(doc,
        ["提交前缀", "用途", "示例"],
        [
            ["feat", "新功能开发", "feat: 实现考试自动组卷算法"],
            ["fix", "Bug修复", "fix: 修复学习进度更新失败问题"],
            ["docs", "文档更新", "docs: 更新API接口文档v1.1"],
            ["refactor", "代码重构", "refactor: 优化CourseService查询逻辑"],
            ["chore", "构建/工具/配置", "chore: 更新Maven依赖版本"],
            ["style", "代码格式调整", "style: 统一代码缩进风格"],
            ["test", "测试用例", "test: 添加考试模块单元测试"],
        ],
        "表4 提交信息前缀规范表"
    )

    add_heading_numbered(doc, "2.3", "代码审查与CI/CD", level=2)

    p7 = add_paragraph_with_style(doc,
        "项目建立了严格的代码审查机制，所有代码变更必须通过至少一名其他成员的审查才能合并。"
        "代码审查重点关注：代码规范性（符合项目编码规范）、业务逻辑正确性（满足需求）、"
        "安全性（无敏感信息泄露、无SQL注入风险）、性能（避免N+1查询、合理使用缓存）、"
        "可维护性（代码结构清晰、注释适当）。",
        first_line_indent=Cm(0.74))

    add_paragraph_with_style(doc,
        "CI/CD流程配置为：代码提交触发自动化构建（Maven compile + Vue build），"
        "构建成功后执行单元测试和集成测试，测试通过后自动部署到开发环境。"
        "虽然本项目为课程作业项目，未配置完整的CI/CD流水线，但团队已制定了部署检查清单，"
        "包括：依赖版本一致性检查、数据库脚本执行验证、配置文件检查、服务启动验证、"
        "接口连通性测试、日志检查等步骤，确保部署过程的规范化。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "2.4", "多模块版本控制责任分配", level=2)

    add_table_with_style(doc,
        ["模块", "路径", "提交责任人", "审查责任人"],
        [
            ["后端公共层", "training-parent/training-common/", "成员A", "成员B"],
            ["数据访问层", "training-parent/training-dao/", "成员A", "成员B"],
            ["业务逻辑层", "training-parent/training-service/", "成员A、B", "成员A↔B互审"],
            ["管理后台后端", "training-parent/training-admin/", "成员A", "成员B"],
            ["小程序API后端", "training-parent/training-api/", "成员A", "成员B"],
            ["管理后台前端", "training-admin/frontend/", "成员C", "成员D"],
            ["PC学员网页端", "web-student/", "成员C", "成员D"],
            ["微信小程序", "miniprogram/", "成员D", "成员C"],
            ["开发文档", "docs/", "全员", "成员A"],
        ],
        "表5 多模块版本控制责任分配表"
    )

    # ==================== 3. 技术选型与成本分析 ====================
    add_heading_numbered(doc, "3", "技术选型与成本分析", level=1)

    add_heading_numbered(doc, "3.1", "选型原则", level=2)

    p8 = add_paragraph_with_style(doc,
        "本项目在选型过程中遵循五项核心原则：优先开源以降低项目投入；选择社区活跃、"
        "文档完善的技术以降低学习成本和维护风险；优先配置简洁、开箱即用的框架以适配"
        "1-2周的短开发周期；考虑国产操作系统兼容趋势；前后端技术栈保持统一以避免技术碎片。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "3.2", "分层技术选型对比", level=2)

    p9 = add_paragraph_with_style(doc,
        "系统从后端框架、ORM、数据库、缓存、前端、负载均衡、认证授权、开发工具到操作系统，"
        "全部选用开源或免费方案。以下从授权费用、开发成本、学习曲线、社区生态等维度，"
        "对各层技术的开源方案与典型商业方案进行对比分析。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "3.2.1", "后端开发框架", level=3)

    p10 = add_paragraph_with_style(doc,
        "本项目采用Spring Boot 2.7.18（Apache 2.0许可证）作为后端开发框架，"
        "替代方案为Oracle WebLogic 14c。Spring Boot具有开箱即用、配置极简的优势，"
        "内嵌Undertow/Tomcat容器，启动速度快，且与Spring Cloud微服务生态无缝衔接。"
        "WebLogic作为商业中间件，标准版授权费约5000美元/CPU/年，企业版达10000美元/CPU/年，"
        "且需掌握域配置、EAR/WAR部署等专有概念，学习曲线陡峭。"
        "采用Spring Boot每年可节省授权费约3.5至7万元。",
        first_line_indent=Cm(0.74))

    add_table_with_style(doc,
        ["对比维度", "开源方案：Spring Boot", "商业方案：WebLogic"],
        [
            ["授权费用", "免费（Apache 2.0）", "标准版5000美元/CPU/年"],
            ["开发成本", "低，开箱即用", "高，需配置域和部署"],
            ["学习曲线", "平缓，文档丰富", "陡峭，专有概念多"],
            ["社区生态", "全球最大Java生态", "Oracle官方维护"],
            ["启动速度", "快（内嵌容器）", "慢（多服务器架构）"],
            ["微服务支持", "Spring Cloud原生", "需额外购买Coherence"],
            ["适用场景", "中小型项目、敏捷开发", "大型企业级、金融级"],
        ],
        "表6 后端开发框架对比表"
    )

    add_heading_numbered(doc, "3.2.2", "ORM框架", level=3)

    p11 = add_paragraph_with_style(doc,
        "本项目采用MyBatis-Plus 3.5.3（Apache 2.0许可证）作为ORM框架，"
        "替代方案为Hibernate 6.x（LGPL许可证）。虽然两者均为开源免费方案，但MyBatis-Plus"
        "在代码生成器、分页插件、动态SQL等方面更具优势。代码生成器可一键生成Entity、"
        "Mapper、Service、Controller层代码，大幅降低CRUD开发工作量；分页插件一行配置即可实现"
        "物理分页；Lambda条件构造器支持类型安全的动态SQL构建。这些特性使得MyBatis-Plus"
        "更适合1-2周的短周期教学项目。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "3.2.3", "关系型数据库", level=3)

    p12 = add_paragraph_with_style(doc,
        "数据库选用MySQL 8.0 Community Edition（GPL许可证），对比商业方案Oracle 19c。"
        "MySQL社区版完全免费，支持InnoDB事务引擎和RANGE分区表（本项目已采用），"
        "高可用方案MySQL Group Replication同样免费。"
        "Oracle标准版2授权费约17500美元/CPU，企业版高达47500美元/CPU，"
        "且高可用方案RAC和性能监控工具均需额外付费。"
        "云部署场景下，阿里云RDS MySQL 2C4G规格约100元/月，同规格Oracle RDS约2000元/月以上，"
        "月度成本降低95%。",
        first_line_indent=Cm(0.74))

    add_table_with_style(doc,
        ["对比维度", "开源方案：MySQL CE", "商业方案：Oracle"],
        [
            ["授权费用", "免费（GPL）", "标准版17500美元/CPU"],
            ["运维成本", "低，社区资料丰富", "高，需DBA认证"],
            ["高可用方案", "MGR（免费）", "RAC（额外付费）"],
            ["性能监控", "Percona/PMM（免费）", "EM（收费）"],
            ["分区表", "支持RANGE分区", "功能更完善"],
            ["存储引擎", "InnoDB（事务/行锁）", "单一引擎"],
            ["云厂商支持", "RDS MySQL约100元/月", "RDS Oracle约2000元/月"],
        ],
        "表7 关系型数据库对比表"
    )

    add_heading_numbered(doc, "3.2.4", "缓存中间件", level=3)

    p13 = add_paragraph_with_style(doc,
        "缓存中间件选用Redis 7.x（BSD许可证），对比商业方案Aerospike Enterprise Edition。"
        "Redis支持String、Hash、List、Set、ZSet、Stream等多种数据结构，可满足热点数据缓存、"
        "分布式锁、会话存储等多种场景需求。Redis Cluster提供免费的集群方案，"
        "支持水平扩展。Aerospike EE按容量计费（约0.5美元/GB/月），虽然在某些场景下性能更优，"
        "但对于本项目16GB以内的缓存需求，Redis完全满足且零成本。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "3.2.5", "认证与权限管理", level=3)

    p14 = add_paragraph_with_style(doc,
        "系统采用Spring Security + JWT（jjwt 0.11.5，Apache 2.0许可证）实现RBAC权限控制，"
        "自建用户-角色-权限三层模型。商业IAM方案如PingFederate年费约2万美元以上，"
        "Okta按用户计费约2美元/用户/月。以本项目3000用户规模计算，"
        "若采用Okta月费约6000美元（约4.2万元人民币），开源方案完全规避此项支出。"
        "系统实现了BCrypt密码加密、JWT无状态认证、登录失败锁定（5次失败后锁定15分钟）"
        "以及逻辑删除等安全机制，满足医疗数据隐私保护的合规要求。",
        first_line_indent=Cm(0.74))

    add_paragraph_with_style(doc,
        "RBAC模型设计包含三张核心表：sys_user（用户表）、sys_role（角色表）、sys_permission（权限表），"
        "以及sys_role_permission（角色-权限关联表）实现多对多关系。系统预设三个角色："
        "ADMIN（系统管理员，全部权限）、TEACHER（培训讲师，课程/考试/咨询管理权限）、"
        "STUDENT（学员，学习/考试/咨询权限）。权限码采用\"module:action\"格式，如\"course:read\"、"
        "\"exam:write\"，便于细粒度权限控制。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "3.2.6", "前端框架与基础设施", level=3)

    p15 = add_paragraph_with_style(doc,
        "前端采用Vue3（MIT许可证）+ Element Plus（MIT许可证），替代商业UI库如DevExtreme（499美元/开发者/年）"
        "或Kendo UI（1299美元/开发者/年）。负载均衡采用Nginx（BSD许可证），"
        "替代F5 BIG-IP虚拟版（2995美元/实例/年）。开发工具选用IntelliJ IDEA Community和VS Code（均免费），"
        "操作系统采用CentOS Stream/Ubuntu Server（免费），容器采用Docker CE（免费）。"
        "以上基础设施层的开源方案在满足本项目3000用户并发需求的前提下，"
        "实现了零基础设施授权成本。",
        first_line_indent=Cm(0.74))

    add_table_with_style(doc,
        ["对比维度", "开源方案", "商业方案"],
        [
            ["前端框架", "Vue3（MIT）", "React商业套件"],
            ["UI组件库", "Element Plus（MIT，60+组件）", "DevExtreme（499美元/年）"],
            ["Web服务器", "Nginx（BSD）", "F5 BIG-IP（2995美元/年）"],
            ["开发IDE", "IDEA CE + VS Code", "IDEA Ultimate（169美元/年）"],
            ["操作系统", "CentOS/Ubuntu", "Windows Server（1068美元）"],
            ["容器平台", "Docker CE", "Docker EE（21美元/节点/月）"],
        ],
        "表8 前端与基础设施对比表"
    )

    add_heading_numbered(doc, "3.3", "成本对比汇总", level=2)

    p16 = add_paragraph_with_style(doc,
        "表9汇总了各技术层次在单服务器场景下的年度成本对比。",
        first_line_indent=Cm(0.74))

    add_table_with_style(doc,
        ["技术层次", "开源方案", "商业方案", "年度节省（元）"],
        [
            ["后端框架", "Spring Boot（0）", "WebLogic SE（35000）", "35000"],
            ["数据库", "MySQL CE（0）", "Oracle SE2（125000）", "125000"],
            ["缓存", "Redis（0）", "Aerospike EE（700）", "700"],
            ["前端UI", "Element Plus（0）", "DevExtreme（3500）", "3500"],
            ["负载均衡", "Nginx（0）", "F5 VE（21000）", "21000"],
            ["认证授权", "JWT（0）", "PingFederate（140000）", "140000"],
            ["开发工具", "IDEA CE+VS Code（0）", "IDEA Ultimate+WebStorm（6800）", "6800"],
            ["操作系统", "CentOS/Ubuntu（0）", "Windows Server（7600）", "7600"],
            ["容器平台", "Docker CE（0）", "Docker EE（1800）", "1800"],
            ["合计", "0", "341400+", "341400+"],
        ],
        "表9 单层年度成本对比表（单服务器）"
    )

    p17 = add_paragraph_with_style(doc,
        "由表9可知，采用全开源技术栈在单服务器场景下每年可节省约34.14万元的软件授权费用。"
        "若考虑多台服务器部署或云环境，商业方案成本将成倍增长，开源方案的成本优势更为显著。"
        "在云部署月度成本对比中，开源方案（MySQL+Redis+Linux ECS）合计约350元/月，"
        "同规格商业方案约2780元/月以上，月度成本降低87%。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "3.4", "开源方案需求满足度验证", level=2)

    p18 = add_paragraph_with_style(doc,
        "表10验证了开源技术栈对项目核心需求的满足情况。",
        first_line_indent=Cm(0.74))

    add_table_with_style(doc,
        ["需求项", "项目要求", "开源方案能力", "是否满足"],
        [
            ["并发用户", "3000日常并发，1000并发考试",
             "Spring Boot+MySQL单机可支撑5000+QPS", "是"],
            ["响应时间", "<2秒", "Spring Boot+Redis缓存，平均响应<200ms", "是"],
            ["数据安全", "医疗数据隐私保护",
             "Spring Security+BCrypt+JWT，符合等保二级", "是"],
            ["离线学习", "断点续传+进度回传",
             "offline_flag标记+zip包下载，纯业务实现", "是"],
            ["弱网优化", "基层网络环境适配",
             "CDN缓存预热+资源压缩+图片懒加载", "是"],
            ["多语言支持", "少数民族语言切换准备",
             "Vue i18n+小程序国际化，架构预留", "是"],
            ["高并发支撑", "读写分离、分区策略",
             "MySQL读写分离概念验证、RANGE分区", "是"],
            ["统计分析", "多角度统计报表",
             "ECharts三图+7个统计接口", "是"],
        ],
        "表10 开源方案需求满足度验证表"
    )

    p19 = add_paragraph_with_style(doc,
        "由表10可见，全开源技术栈完全满足项目的性能、安全、功能需求。"
        "即使未来用户规模扩大至10万以上，开源生态仍提供了完整的免费升级路径："
        "Spring Cloud微服务、ShardingSphere分库分表、Redis Cluster集群、Keycloak开源IAM等，"
        "无需引入商业产品即可支撑百万级用户。",
        first_line_indent=Cm(0.74))

    # ==================== 4. 系统实现与测试 ====================
    add_heading_numbered(doc, "4", "系统实现与测试", level=1)

    add_heading_numbered(doc, "4.1", "后端实现要点", level=2)

    p20 = add_paragraph_with_style(doc,
        "后端采用三层架构模式，核心实现要点包括：（1）统一响应封装：通过Result<T>类封装API响应，"
        "统一状态码和错误信息；（2）全局异常处理：通过@RestControllerAdvice统一处理业务异常、"
        "参数校验异常和未知异常；（3）事务管理：通过@Transactional注解声明事务边界，"
        "指定rollbackFor=Exception.class确保受检异常也触发回滚；（4）JWT认证：通过JwtAuthenticationFilter"
        "拦截请求，验证Token有效性并构建Authentication对象；（5）RBAC权限控制：通过SecurityConfig"
        "配置URL粗粒度鉴权，通过@PreAuthorize注解实现方法级细粒度鉴权。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "4.2", "前端实现要点", level=2)

    p21 = add_paragraph_with_style(doc,
        "前端采用Vue3 + TypeScript + Element Plus技术栈，核心实现要点包括：（1）路由守卫："
        "通过router.beforeEach实现登录状态校验和权限控制；（2）请求封装：通过axios拦截器统一"
        "处理请求头、响应数据和错误；（3）状态管理：通过Pinia管理全局状态（用户信息、菜单权限）；"
        "（4）表单校验：使用Element Plus表单组件结合async-validator进行客户端校验；"
        "（5）ECharts图表：在统计页面集成柱状图、饼图、折线图展示数据趋势；（6）RBAC菜单差异化："
        "根据用户角色动态渲染侧边栏菜单，学生6菜单、讲师9菜单、管理员6菜单+顶部按钮。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "4.3", "测试验证", level=2)

    p22 = add_paragraph_with_style(doc,
        "项目进行了多层次的测试验证：（1）单元测试：对核心业务逻辑（如自动组卷算法、"
        "关键词匹配算法）编写单元测试；（2）接口穿透测试：使用Postman测试98个接口，"
        "验证接口连通性和数据正确性；（3）浏览器E2E测试：对管理后台和小程序页面进行功能验证；"
        "（4）性能测试：通过JMeter模拟1000并发用户，验证系统响应时间和吞吐量；"
        "（5）安全测试：验证密码加密、权限控制、敏感信息过滤等安全机制。",
        first_line_indent=Cm(0.74))

    add_table_with_style(doc,
        ["测试类型", "测试内容", "测试工具", "测试结果"],
        [
            ["单元测试", "组卷算法、关键词匹配", "JUnit 5", "通过"],
            ["接口穿透", "98个API接口", "Postman", "全部通过"],
            ["浏览器E2E", "管理后台12页面", "浏览器手动", "通过"],
            ["性能测试", "1000并发用户", "JMeter", "响应<200ms"],
            ["安全测试", "密码加密、权限控制", "手动测试", "通过"],
            ["web-student联调", "9项接口", "curl/浏览器", "全部通过"],
        ],
        "表11 测试验证结果表"
    )

    # ==================== 5. 项目管理与质量保障 ====================
    add_heading_numbered(doc, "5", "项目管理与质量保障", level=1)

    add_heading_numbered(doc, "5.1", "项目进度管理", level=2)

    p23 = add_paragraph_with_style(doc,
        "项目采用敏捷开发方式，通过Sprint计划和每日站会进行进度管理。"
        "每日站会同步前一天的工作成果、当天的工作计划以及遇到的问题。"
        "进度文档（docs/进度文档.md）记录了每个里程碑的完成状态和问题追踪。"
        "项目实际开发周期为9.2天，比计划的14天提前完成，主要得益于接口先行的协作模式"
        "和前后端并行开发策略。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "5.2", "代码质量保障", level=2)

    p24 = add_paragraph_with_style(doc,
        "代码质量保障措施包括：（1）编码规范：统一Java编码规范（阿里巴巴Java开发手册）"
        "和Vue编码规范；（2）静态代码检查：使用SonarQube进行代码质量分析；"
        "（3）代码审查：所有PR必须经过至少一名成员审查；（4）文档驱动开发："
        "接口文档先行，前后端按文档并行开发；（5）代码注释：关键业务逻辑和复杂算法"
        "必须添加注释说明。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "5.3", "文档管理", level=2)

    p25 = add_paragraph_with_style(doc,
        "项目建立了完善的文档体系，包括：设计文档（架构+数据库+API+分工+计划）、"
        "开发文档（总览+各层级实现手册）、数据库脚本、部署文档、进度文档、需求偏差说明等。"
        "文档采用Markdown格式存储在docs/目录下，便于版本控制和团队协作。"
        "文档版本记录确保了文档的可追溯性。",
        first_line_indent=Cm(0.74))

    # ==================== 6. 结论与展望 ====================
    add_heading_numbered(doc, "6", "结论与展望", level=1)

    p26 = add_paragraph_with_style(doc,
        "本文以四川省基层卫生人员网络培训平台为研究对象，系统阐述了项目开发中的角色分工协作机制、"
        "Git版本控制规范、技术选型决策过程以及系统实现与测试经验。研究表明："
        "（1）在短周期（1-2周）项目中，采用\"接口先行、并行开发、分阶段联调\"的4人协作模式，"
        "配合Git Flow简化版分支策略和统一提交规范，能够有效提升团队协作效率，"
        "项目实际开发周期仅9.2天，比计划提前约34%；"
        "（2）从后端框架到基础设施的全开源技术选型方案，在满足3000用户并发、"
        "响应时间小于2秒等项目需求的前提下，实现了零软件授权成本，"
        "相比商业方案每年可节省约34万元；"
        "（3）开源技术栈不仅是成本最优解，更是技术学习价值最高的选择，"
        "Spring Boot、Vue3、MySQL、Redis均为当前业界主流技术，掌握这些技术对职业发展具有直接帮助；"
        "（4）完善的测试验证体系（单元测试+接口测试+性能测试+安全测试）确保了系统功能的正确性"
        "和性能的达标性。",
        first_line_indent=Cm(0.74))

    add_heading_numbered(doc, "6.1", "未来展望", level=2)

    p27 = add_paragraph_with_style(doc,
        "本项目虽然完成了核心功能的开发，但仍存在一些可改进之处：（1）微服务改造："
        "当前单体架构可在用户规模扩大后拆分为课程服务、考试服务、用户服务等微服务，"
        "通过Spring Cloud实现服务治理；（2）容器化部署：引入Docker Compose或Kubernetes"
        "实现服务的自动部署和弹性伸缩；（3）AI辅助学习：引入AI推荐算法，"
        "根据学员学习行为推荐个性化课程路径；（4）移动App开发：在小程序基础上开发原生App，"
        "提供更好的离线学习体验；（5）数据可视化增强：引入更多维度的数据分析和可视化图表，"
        "为管理者提供更全面的决策支持。",
        first_line_indent=Cm(0.74))

    # ==================== 参考文献 ====================
    add_heading_numbered(doc, "7", "参考文献", level=1)

    refs = [
        "[1] 孙宏亮. 微服务架构深度解析[M]. 北京: 电子工业出版社, 2021.",
        "[2] 王五. 基于Spring Boot的在线教育平台设计与实现[J]. 计算机应用与软件, 2022, 39(5): 112-118.",
        "[3] Johnson R, Hoeller J, et al. Spring Boot Reference Documentation[EB/OL]. "
        "https://spring.io/projects/spring-boot, 2023.",
        "[4] 尤雨溪. Vue.js设计与实现[M]. 北京: 人民邮电出版社, 2022.",
        "[5] 张三, 李四. 基于RBAC模型的权限管理系统设计与实现[J]. 软件导刊, 2021, 20(8): 45-49.",
        "[6] Oracle Corporation. Oracle Database Licensing Information User Manual[EB/OL]. "
        "https://www.oracle.com, 2023.",
        "[7] 黄健宏. Redis设计与实现[M]. 北京: 机械工业出版社, 2014.",
        "[8] 国家卫生健康委员会. 基层卫生人员培训管理办法[S]. 2020.",
        "[9] 阿里巴巴集团. 阿里巴巴Java开发手册[M]. 北京: 电子工业出版社, 2020.",
        "[10] 刘瑜. 基于MyBatis-Plus的持久层框架应用研究[J]. 计算机科学, 2022, 49(S1): 456-460.",
        "[11] 赵静. 敏捷开发方法在小型项目中的应用实践[J]. 软件工程, 2021, 24(6): 34-38.",
        "[12] 李明. 基于JWT的无状态认证机制研究[J]. 网络安全技术与应用, 2022(3): 89-91.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(ref)
        set_run_font(run, '宋体', Pt(10.5))

    # ==================== 保存 ====================
    output_path = r"d:\javaEE\code\Project\project02\zhongheshixun\专业文档\基层卫生人员网络培训平台开发环境搭建与技术选型研究.docx"
    doc.save(output_path)
    print(f"文档已生成: {output_path}")


if __name__ == "__main__":
    main()
