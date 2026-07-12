"""
从 Markdown 源重新生成学生甲可行性分析报告 docx，应用统一样式：
- A4 页边距
- 中文字体：宋体/微软雅黑，英文字体：Times New Roman
- 标题 1/2/3 层级（黑体加粗）
- 表头灰色底纹 + 全边框 + 自动列宽
- 页眉：文档标题；页脚：页码（居中）
- 代码块/流程图：等宽字体 + 浅灰底纹 + 边框
- 引用块：左侧竖线 + 浅蓝底纹
- 生成目录域（Word 打开后更新域即可）
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


INPUT_MD = (
    "D:/A-Users/Desktop/zhongheshixun/defense-materials/"
    "综合实训可行性分析/A-学生甲-项目可行性分析报告.md"
)
OUTPUT_DOCX = (
    "D:/A-Users/Desktop/zhongheshixun/defense-materials/"
    "综合实训可行性分析/A-学生甲-项目可行性分析报告.docx"
)

DOC_TITLE = "项目可行性分析报告"

# ---------- 字体配置 ----------
CN_BODY = "宋体"          # 正文中文
CN_HEADING = "黑体"       # 标题中文
EN_FONT = "Times New Roman"
CODE_FONT = "Consolas"    # 代码块字体

# ---------- 字号 ----------
SIZE_TITLE = Pt(22)
SIZE_H1 = Pt(18)
SIZE_H2 = Pt(15)
SIZE_H3 = Pt(13)
SIZE_BODY = Pt(12)
SIZE_CODE = Pt(10.5)
SIZE_META = Pt(10.5)


def set_cn_font(run, cn_font=CN_BODY, en_font=EN_FONT, size=SIZE_BODY, bold=False):
    """设置 run 的中英文字体。"""
    run.font.name = en_font
    run.font.size = size
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    r.rPr.rFonts.set(qn("w:ascii"), en_font)
    r.rPr.rFonts.set(qn("w:hAnsi"), en_font)


def set_paragraph_spacing(p, before=0, after=2, line=1.5):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12)
    run = p.add_run(text)
    set_cn_font(run, cn_font=CN_HEADING, size=SIZE_TITLE, bold=True)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_cn_font(run, cn_font=CN_HEADING, size=SIZE_H1, bold=True)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6B)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_cn_font(run, cn_font=CN_HEADING, size=SIZE_H2, bold=True)
    run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_cn_font(run, cn_font=CN_HEADING, size=SIZE_H3, bold=True)
    return p


def add_body(doc, text):
    """常规段落；保留行内加粗/代码。"""
    p = doc._body_p_add_next() if False else doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line=1.5)
    # 简单处理行内 **bold** 和 `code`
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_cn_font(run, cn_font=CODE_FONT, en_font=CODE_FONT, size=SIZE_BODY)
            run.font.color.rgb = RGBColor(0xA8, 0x00, 0x00)
            rPr = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "FFF0F0")
            rPr.append(shd)
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_cn_font(run, size=SIZE_BODY, bold=True)
        else:
            run = p.add_run(part)
            set_cn_font(run, size=SIZE_BODY)
    return p


def _add_toc(doc):
    """插入目录域（Word 中右键更新域即可）。"""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar)

    run2 = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._element.append(fldChar2)

    # 占位文本（更新域后会被替换）
    run4 = p.add_run("（请在 Word 中右键「更新域」以生成目录）")
    set_cn_font(run4, size=SIZE_BODY)
    run4.italic = True
    run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    run5 = p.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._element.append(fldChar3)
    return p


def set_cell_borders(cell, color="BFBFBF"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table(doc, header, rows):
    """添加带样式表格：表头灰色底纹加粗，全边框，正文宋体。"""
    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 全表格默认字体
    table.style = "Table Grid"
    # 表头
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_cn_font(run, cn_font=CN_HEADING, size=SIZE_BODY, bold=True)
        shade_cell(cell, "D9D9D9")
        set_cell_borders(cell)
    # 数据行
    for r_idx, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            cell = table.rows[r_idx].cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_cn_font(run, size=SIZE_BODY)
            set_cell_borders(cell)
            if r_idx % 2 == 0:
                shade_cell(cell, "F2F2F2")
    # 列宽自动
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(2.5)
    return table


def add_code_block(doc, lines):
    """代码块/流程图：浅灰底纹 + 等宽字体 + 边框。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F5F5F5")
    set_cell_borders(cell, color="CCCCCC")
    cell.text = ""
    for line in lines:
        p = cell.paragraphs[0] if cell.paragraphs[0].text == "" else cell.add_paragraph()
        run = p.add_run(line)
        set_cn_font(run, cn_font=CODE_FONT, en_font=CODE_FONT, size=SIZE_CODE)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
    return table


def add_blockquote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, "EEF4FB")
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), "2E5C8A")
    tcBorders.append(left)
    tcPr.append(tcBorders)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_cn_font(run, size=SIZE_BODY)
    run.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return table


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ==================== 读取 Markdown ====================
with open(INPUT_MD, "r", encoding="utf-8") as f:
    md = f.read()

lines = md.split("\n")
total = len(lines)


# ==================== 构建 Word ====================
doc = Document()

# 页面设置
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 默认正文样式
style = doc.styles["Normal"]
style.font.name = EN_FONT
style.font.size = SIZE_BODY
style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_BODY)

# ---------- 页眉页脚 ----------
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run(DOC_TITLE)
set_cn_font(hr, cn_font=CN_HEADING, size=Pt(10))
hr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
# 页眉下横线
pPr = hp._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:color"), "BFBFBF")
pBdr.append(bottom)
pPr.append(pBdr)

footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# PAGE 域
run = fp.add_run()
fldChar = OxmlElement("w:fldChar")
fldChar.set(qn("w:fldCharType"), "begin")
run._element.append(fldChar)
run2 = fp.add_run()
instr = OxmlElement("w:instrText")
instr.text = "PAGE   \\* MERGEFORMAT"
run2._element.append(instr)
run3 = fp.add_run()
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "end")
run3._element.append(fldChar2)
run_label = fp.add_run(" 页")
set_cn_font(run_label, size=Pt(10))
run_label.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


# ---------- 逐行渲染 ----------
i = 0


def strip_inline(text):
    """清空行内加粗/代码符号，返回纯文本。"""
    return re.sub(r"\*+|\`", "", text)


def slugify(text):
    s = re.sub(r"\s+", "", text)
    s = re.sub(r"[^\w一-鿿\-]", "", s, flags=re.UNICODE)
    return s.lower()


heading_slugs = []  # (level, text) 用于手动目录


while i < total:
    raw = lines[i]
    line = raw.rstrip("\n")

    # 标题（检测前置 # 行）
    m_h = re.match(r"^(#{1,6})\s+(.*)$", line)
    if m_h:
        level = len(m_h.group(1))
        text = m_h.group(2).strip()
        text = re.sub(r"\s*\{#.*\}$", "", text)  # 去掉 {#...}
        if level == 1 and i == 0:
            add_title(doc, text)
        elif level == 1:
            add_title(doc, text)
        elif level == 2:
            add_h1(doc, text)
            heading_slugs.append((2, text))
        elif level == 3:
            add_h2(doc, text)
            heading_slugs.append((3, text))
        else:
            add_h3(doc, text)
            heading_slugs.append((4, text))
        i += 1
        continue

    # 目录占位
    if line.strip() == "## 目录":
        add_h1(doc, "目录")
        _add_toc(doc)
        i += 1
        continue

    # 分隔线
    if re.match(r"^-{3,}\s*$", line.strip()):
        add_horizontal_line(doc)
        i += 1
        continue

    # 代码块
    if line.strip().startswith("```"):
        code_lines = []
        i += 1
        while i < total and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i].rstrip("\n"))
            i += 1
        i += 1  # 结束 ```
        if code_lines:
            add_code_block(doc, code_lines)
        continue

    # 表格
    if "|" in line and re.match(r"^\s*\|", line):
        # 收集连续表格行
        tbl_lines = []
        while i < total and "|" in lines[i] and re.match(r"^\s*\|", lines[i]):
            tbl_lines.append(lines[i])
            i += 1
        # 解析 header | separator | rows
        header_line = tbl_lines[0]
        header = [c.strip() for c in re.split(r"(?<!\\)\|", header_line.strip("| "))]
        rows = []
        for r in tbl_lines[2:]:
            if re.match(r"^\s*\|[\s\-:|]+\|", r):
                continue
            row = [c.strip() for c in re.split(r"(?<!\\)\|", r.strip("| "))]
            rows.append(row)
        if rows:
            add_table(doc, header, rows)
        continue

    # 引用块
    if line.strip().startswith(">"):
        q_lines = []
        while i < total and lines[i].strip().startswith(">"):
            q_lines.append(re.sub(r"^>\s?", "", lines[i]))
            i += 1
        add_blockquote(doc, "\n".join(q_lines).strip())
        continue

    # 无序列表
    if re.match(r"^\s*[-*]\s+", line):
        items = []
        while i < total and re.match(r"^\s*[-*]\s+", lines[i]):
            items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]))
            i += 1
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            for part in re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", item):
                if part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    set_cn_font(run, cn_font=CODE_FONT, en_font=CODE_FONT, size=SIZE_BODY)
                    run.font.color.rgb = RGBColor(0xA8, 0x00, 0x00)
                elif part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_cn_font(run, size=SIZE_BODY, bold=True)
                else:
                    run = p.add_run(part)
                    set_cn_font(run, size=SIZE_BODY)
        continue

    # 有序列表
    if re.match(r"^\s*\d+[\.\)]\s+", line):
        items = []
        while i < total and re.match(r"^\s*\d+[\.\)]\s+", lines[i]):
            items.append(re.sub(r"^\s*\d+[\.\)]\s+", "", lines[i]))
            i += 1
        for item in items:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            for part in re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", item):
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_cn_font(run, size=SIZE_BODY, bold=True)
                else:
                    run = p.add_run(part)
                    set_cn_font(run, size=SIZE_BODY)
        continue

    # 空行
    if line.strip() == "":
        i += 1
        continue

    # 普通段落（带行内格式）
    add_body(doc, line.strip())
    i += 1


doc.save(OUTPUT_DOCX)
print("Done: " + OUTPUT_DOCX)
print("Headings: " + str(len(heading_slugs)))
